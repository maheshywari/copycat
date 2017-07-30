#code to copy the text being highlighted while reading documents and provide them in a pdf format

import os
import sys
import win32api
import tkinter as tk
#get the path from where script is being executed
script_path=os.path.realpath('copycat.py')
split_path=os.path.split(script_path)
base_path=split_path[0]

state_left = win32api.GetKeyState(0x01)  # Left button down = 0 or 1. Button up = -127 or -128
state_right = win32api.GetKeyState(0x02)  # Right button down = 0 or 1. Button up = -127 or -128
import tkinter as tk
import PIL.Image
from PIL import ImageTk
from itertools import count
import win32com.client
import subprocess
import time
import threading
import queue
import win32api
import win32con
import signal
import pyautogui
import win32clipboard
from docx.shared import Pt
from docx.enum.text import *
from docx.enum.style import *
from docx import Document
from tkinter import ttk
from tkinter import *
from array import *
from tkinter import filedialog
if os.path.exists(os.path.join(base_path,"exception_log.txt")):
	os.remove(os.path.join(base_path,"exception_log.txt"))
exception_log=open("exception_log.txt","a")
document = Document()
document_h=Document()
document_s=Document()

#global varaibles for copying
text=[]
text_final=[]
heading_index=[]
sub_heading_index=[]
heading=[]
sub_heading=[]
loop_1=0
loop=0
data_dummy=""
data=""
text_body=[]
exception_array=[]
app2=""
listen_key_flag=[]
alert_flag=""

#global variables for animation
frames = []
loc = 0
delay=0

#global variables for output files
folder_path=""
document_name=""
docu_format=""

#function to create document
def copy2_word(parameter,content):
	global document
	if parameter == "heading":
		#font.name = 'Calibri'
		#font.underline = True
		#font.size = Pt(14)
		#font.bold=True
		document.add_heading(content,0)
	elif parameter == "sub-heading":
		#font.name = 'Calibri'
		#font.underline = WD_UNDERLINE.DOTTED
		#font.size = Pt(12)
		#font.bold=True
		document.add_heading(content,1)
	else:
		#font.name = 'Calibri'
		#font.size = Pt(10)
		document.add_paragraph(content, style='ListBullet')
	document.save('copy_cat.docx')

#function to collect data
def process_text(text_1):
	global text_final,heading_index,sub_heading_index
	text_dummy=text_1
	for m_c in range (0,len(text_dummy)):
		entered_loop=False
		for h_c in heading_index:
			if m_c== h_c:
				text_head=text_dummy[m_c]
				copy2_word("heading",text_head)
				entered_loop=True
				break
		for sh_c in sub_heading_index:
			if m_c == sh_c:
				text_sub_head=text_dummy[m_c]
				copy2_word("sub-heading",text_sub_head)
				entered_loop=True
		if entered_loop!=True:
			text_body=text_dummy[m_c]
			copy2_word("text_body",text_body)

#thread for the options
class options(threading.Thread):
	global data
	def __init__(self):
		threading.Thread.__init__(self)
		self.start()
	
	def callback(self):
		self.root.quit()	
	
	def run(self):
		self.root= tk.Tk()
		self.root.protocol("WM_DELETE_WINDOW", self.callback)
		self.root.title("OPTIONS")
		mainframe=Frame(self.root,width=150,height=200).pack()
		self.canvas=Canvas(mainframe,bg='white')
		self.canvas.place(anchor=NW,width=150,height=200)
		self.root.call('wm', 'attributes', '.', '-topmost', '1')
		button_text=["HEADING","SUB-HEADING","UNDO"]
		self.button=[]
		self.step=40
		for i in range (0,3):
			self.button.append(Button(self.canvas,text=button_text[i],width=15,fg="black"))
			self.button[i].pack()
			self.canvas.create_window(80,self.step,window=self.button[i])
			self.step=self.step+40
		self.button[0].configure(command=lambda:self.heading_fc())
		self.button[1].configure(command=lambda:self.sub_heading_fc())
		self.button[2].configure(command=lambda:self.undo_fc())
		self.root.mainloop()
	
	def heading_fc(self):
		global data,heading_index,text,heading,listen_key_flag
		#print ("heading")
		index=len(text)
		heading_index.append(index-1)
		heading.append(text[index-1])
		listen_key_flag.append("heading")
		
	def sub_heading_fc(self):
		global data,sub_heading_index,text,sub_heading,listen_key_flag
		index=len(text)
		sub_heading_index.append(index-1)
		sub_heading.append(text[index-1])
		listen_key_flag.append("sub_heading")
		
	def undo_fc(self):
		global data,heading_index,text,heading,listen_key_flag,sub_heading_index,alert_flag
		print ("undoing")
		print (len(listen_key_flag))
		for cnt in range (0,len(listen_key_flag)):
			if listen_key_flag[cnt]=="heading":
				for cnt in range (0,len(heading_index)):
					#print (heading[cnt])
					heading_index.pop(-1)
				listen_key_flag.pop(-1)
			elif listen_key_flag[cnt]=="sub_heading":
				for cnt in range (0,len(sub_heading_index)):
					sub_heading_index.pop(-1)
				listen_key_flag.pop(-1)
			elif listen_key_flag[cnt]=="text_data":
				for cnt in range (0,len(text)):
					text.pop(-1)
				listen_key_flag.pop(-1)
		if len(listen_key_flag) <= 0:
			alert_msg=Text(self.root)
			alert_msg.insert(INSERT,"All changes are undone")
			alert_msg.configure(foreground="red")
			#time.sleep(2)
			alert_msg.configure(foreground="black")
			alert_msg.pack()
			alert_flag=True
			#self.root1.mainloop()
		else:
			if alert_flag == True:
				alert_msg.delete(1.0,END)
				alert_flag=False
		
#main program thread			
class copy_thread(threading.Thread):
	global app2,data,heading_index,text,heading
	def __init__(self):
		threading.Thread.__init__(self)
		self.start()
	def run(self):
		#print ("i am in ")
		while True:
			if (app2.isAlive()):
				#print ("i am in ")
				global text,state_left,heading,sub_heading,document,document_h,document_s,exception_array,listen_key_flag
				a = win32api.GetKeyState(0x01)
				b = win32api.GetKeyState(0x02)
				data=""
				if a != state_left:
					state_left = a
					if a < 0:
						dummy_1=0
					else:
						#print('Left Button Released')
						try:
							win32clipboard.OpenClipboard()
							win32clipboard.EmptyClipboard()
							win32clipboard.CloseClipboard()
							pyautogui.hotkey('ctrl', 'c')
							win32clipboard.OpenClipboard()
							if win32clipboard.EnumClipboardFormats() !=0:
								data=win32clipboard.GetClipboardData()
							#print (win32clipboard.EnumClipboardFormats())
								win32clipboard.EmptyClipboard()
								text.append(data)
								listen_key_flag.append("text_data")
						#owner= win32clipboard.GetClipboardOwner()
						#print (win32clipboard.EnumClipboardFormats())
							win32clipboard.CloseClipboard()
						#print (data)
						except Exception as e:
							exception_array.append(e)
						#print (win32clipboard.GetClipboardOwner())
							continue	
			
			else:
				process_text(text)
				quit()

#callback for two threads
def main():
	global app2
	app2=options()
	app1=copy_thread()

class main_frame():
	global root,tk,document_name		
	def __init__(self):
		self.root= tk.Tk()
		#self.root.protocol("WM_DELETE_WINDOW", self.callback)
		self.root.title("COPY_CAT")
		#self.root.wm_iconbitmap('D:\\Scripting\\Scripting_files\\copycat_data\\cat.ico')
		mainframe=Frame(self.root,width=1200,height=750).pack()
		self.canvas=Canvas(mainframe,bg='white')
		self.canvas.place(anchor=NW,width=1200,height=1000)
		self.canvas.create_line(20,450,500,450,fill="black")
		self.canvas.create_line(660,450,1180,450,fill="black")
		canvas_id = self.canvas.create_text(520,430, anchor="nw",font="Cambria 20   ",fill="RED")
		self.canvas.itemconfig(canvas_id, text="SETTINGS")
		canvas_id = self.canvas.create_text(40,20, anchor="nw",font="Cambria 16   ",fill="magenta")
		self.canvas.itemconfig(canvas_id, text="COPY_CAT")
		self.knock=PhotoImage(file = 'D:\\Scripting\\Scripting_files\\copycat_data\\images\\knock.png')
		self.canvas.create_image(190, 34, image=self.knock)
		self.step=90
		self.textvaraible=["Hi All,","This is an small GUI developed by me.It will help you to prepare notes","while reading newspaper or bla bla bla"]
		for i in range(0,3):
			canvas_id = self.canvas.create_text(40,self.step, anchor="nw",font="Cambria 14   ")
			self.canvas.itemconfig(canvas_id, text=self.textvaraible[i])
			self.step=self.step+25
		self.eee=PhotoImage(file = 'D:\\Scripting\\Scripting_files\\copycat_data\\images\\eee.png')
		self.canvas.create_image(370, 148, image=self.eee)
		self.textvaraible=["How to Use:","->Open a newspaper in any search engine","->Drag the left mouse over the data which you want to select"]
		self.textvaraible.append("->Once you release the mouse the data will get copied")
		self.textvaraible.append("->For heading select the text and press h")
		self.textvaraible.append("->For sub-heading select the text and press s")
		self.step=190
		for i in range(0,6):
			canvas_id = self.canvas.create_text(40,self.step, anchor="nw",font="Cambria 14   ")
			self.canvas.itemconfig(canvas_id, text=self.textvaraible[i])
			self.step=self.step+25
			
		#document path
		canvas_id = self.canvas.create_text(20,500, anchor="nw",font="Cambria 12   ",fill="blue")
		self.canvas.itemconfig(canvas_id, text="Document path : ")
		self.entry=Entry(self.canvas,width=50,bd=4)
		self.canvas.create_window(430,510,window=self.entry)
		self.button=Button(self.canvas,text="Browse",command=lambda:self.openfile(),width=10,height=1)
		self.canvas.create_window(650,509,window=self.button)
		
		#document name
		canvas_id = self.canvas.create_text(20,550, anchor="nw",font="Cambria 12   ",fill="blue")
		self.canvas.itemconfig(canvas_id, text="Document Name : ")
		self.entry_name=Entry(self.canvas,width=30,bd=4)
		self.canvas.create_window(380,560,window=self.entry_name)
		document_name=self.entry_name.get()
		
		#document format
		canvas_id = self.canvas.create_text(20,610, anchor="nw",font="Cambria 12   ",fill="blue")
		self.canvas.itemconfig(canvas_id, text="Document format : ")
		self.textvaraible=["PDF","DOC"]
		self.step1=300
		self.step2=330
		self.Checkbutton2=[]
		for i in range(0,2):
			self.Checkbutton2.append(Checkbutton(self.canvas,offvalue=0,onvalue=1))
			self.Checkbutton2[i].pack()
			self.Checkbutton2[i].deselect()
			self.canvas.create_window(self.step1,620,window=self.Checkbutton2[i])
			canvas_id = self.canvas.create_text(self.step2,610, anchor="nw",font="CALIBRI 11 ")
			self.canvas.itemconfig(canvas_id,text=self.textvaraible[i])
			self.step1=self.step1+90
			self.step2=self.step2+90
		self.form1=IntVar()
		self.form2=IntVar()
		self.Checkbutton2[0].config(variable=self.form1,command=lambda:self.sel(self.form1))
		self.Checkbutton2[1].config(variable=self.form2,command=lambda:self.sel(self.form2))
			
		#command options
		#queue_1 = queue.queue()
		self.button=Button(self.canvas,text="START THE MISSION",width=20,fg="black",command=lambda:self.call())
		self.button.pack()
		self.canvas.create_window(600,700,window=self.button)
		
		#animation
		self.animagif=Label(self.root)
		self.animagif.pack()
		#photo = tk.PhotoImage(file ='D:\\Scripting\\Scripting_files\\copycat_data\\cat.gif')
		self.canvas.create_window(930,200,window=self.animagif)
		#self.load('D:\\Scripting\\Scripting_files\\copycat_data\\cat.gif')
		#self.root_1 = tk.Tk()
		self.img='D:\\Scripting\\Scripting_files\\copycat_data\\gifs\\animated_cat_on_book.gif'
		self.root.after(0, self.load, self.img)
		#self.root_1.mainloop()
		self.root.mainloop()
	
	def call(self):
		self.root.destroy()
		main()
		
	#Function to select dodument format	
	def sel(self,text):
		global docu_format
		if self.form1.get() == 1:
			docu_format="PDF"
		if self.form2.get() == 1:
			docu_format="DOC"
		#print (docu_format)
			
	#directory search		
	def openfile(self):
		global folder_path
		folder_path=(filedialog.askdirectory())
		self.entry.delete(0,END)
		self.entry.insert(1,folder_path)
		
	#load animation
	def load(self,im):
		global label,delay
		#print (im)
		if isinstance(im, str):
			im = PIL.Image.open(im)
		try:
			for i in count(1):
				frames.append(ImageTk.PhotoImage(im.copy()))
				im.seek(i)
		except EOFError:
			pass
			
		try:
			delay = im.info['duration']
		except:
			delay = 100
			
		if len(frames) == 1:
			config(image=frames[0])
		else:
			self.next_frame()
		
	def unload(self):
		config(image=None)
		frames = None
		
	def next_frame(self):
		global loc,label,delay
		#print (loc)
		if frames:
			loc += 1
			loc %= len(frames)
			self.animagif.config(image=frames[loc])
			self.root.after(delay, self.next_frame)
			
app = main_frame()